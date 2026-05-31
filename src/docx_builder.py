"""
venturecompany 디자인 시스템 기반 .docx 생성기
design.md의 컴포넌트를 정확히 구현
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ================ 색상 시스템 (design.md §2) ================
NAVY        = RGBColor(0x05, 0x1C, 0x2A)
BLUE        = RGBColor(0x16, 0x3E, 0x93)
CYAN        = RGBColor(0x30, 0xA3, 0xDA)
BLACK       = RGBColor(0x06, 0x02, 0x00)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = "F2F2F2"
GRAY_BORDER = "D9D9D9"
NAVY_HEX    = "051C2A"
BLUE_HEX    = "163E93"
CYAN_HEX    = "30A3DA"

FONT = "Arial"


# ================ 헬퍼: XML 직접 조작 ================
def _shade_cell(cell, hex_color):
    """셀 배경색"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def _set_cell_borders(cell, color="D9D9D9", size="4"):
    """셀 테두리"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

def _bottom_border_paragraph(paragraph, color_hex, size="18"):
    """단락 하단에 굵은 색 라인 (액센트 바)"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def _top_border_paragraph(paragraph, color_hex, size="18"):
    """단락 상단 라인 (콜아웃 박스용)"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), size)
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color_hex)
    p_bdr.append(top)
    p_pr.append(p_bdr)

def _set_char_spacing(run, value):
    """글자 자간 (1/20 pt 단위)"""
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(value))
    rPr.append(spacing)


# ================ 컴포넌트 함수 ================
def add_run(paragraph, text, *, size=10.5, bold=False, color=BLACK, font=FONT, char_spacing=None):
    run = paragraph.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    # 한글 폰트도 같이 지정
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    if char_spacing is not None:
        _set_char_spacing(run, char_spacing)
    return run

def eyebrow(doc, text):
    """이브로우 (CYAN 8pt 대문자, 자간 +40)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text.upper(), size=8, bold=True, color=CYAN, char_spacing=40)
    return p

def h1(doc, text):
    """H1: NAVY 18pt Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=18, bold=True, color=NAVY)
    return p

def h2(doc, text):
    """H2: BLUE 13pt Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=13, bold=True, color=BLUE)
    return p

def h3(doc, text):
    """H3: NAVY 11pt Bold"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=11, bold=True, color=NAVY)
    return p

def body(doc, text):
    """본문: BLACK 10.5pt, justified, 1.25 line spacing"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, size=10.5, color=BLACK)
    return p

def accent_bar(doc, color_hex=CYAN_HEX, size="18"):
    """액센트 바 (CYAN 라인)"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    _bottom_border_paragraph(p, color_hex, size=size)
    return p

def bullet(doc, text, level=0):
    """불릿: level 0 = ■ CYAN, level 1 = – BLUE"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    if level == 0:
        add_run(p, "■  ", size=10.5, color=CYAN, bold=True)
    else:
        add_run(p, "–  ", size=10.5, color=BLUE, bold=True)
    add_run(p, text, size=10.5, color=BLACK)
    return p

def lead_in_bullet(doc, lead, rest):
    """Lead-in 불릿: ■ 굵은 NAVY lead — 일반 BLACK rest"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "■  ", size=10.5, color=CYAN, bold=True)
    add_run(p, f"{lead} — ", size=10.5, color=NAVY, bold=True)
    add_run(p, rest, size=10.5, color=BLACK)
    return p

def numbered(doc, items):
    """번호 리스트: BLUE bold marker"""
    for idx, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(3)
        add_run(p, f"{idx}.  ", size=10.5, color=BLUE, bold=True)
        add_run(p, str(text), size=10.5, color=BLACK)

def kpi_strip(doc, items):
    """KPI 스트립: 3개 타일, NAVY/BLUE/CYAN 배경"""
    if len(items) != 3:
        # 3개가 아니면 일반 표로 fallback
        items = (items + [{"value": "-", "label": "-"}] * 3)[:3]

    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    bg_colors = [NAVY_HEX, BLUE_HEX, CYAN_HEX]

    for col, (item, bg) in enumerate(zip(items, bg_colors)):
        # 값 셀
        val_cell = table.cell(0, col)
        val_cell.width = Cm(5.5)
        _shade_cell(val_cell, bg)
        _set_cell_borders(val_cell, color="FFFFFF", size="12")
        val_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = val_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(8)
        p1.paragraph_format.space_after = Pt(2)
        add_run(p1, str(item.get('value', '-')), size=22, bold=True, color=WHITE)

        # 라벨 셀
        lab_cell = table.cell(1, col)
        lab_cell.width = Cm(5.5)
        _shade_cell(lab_cell, bg)
        _set_cell_borders(lab_cell, color="FFFFFF", size="12")
        lab_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p2 = lab_cell.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(8)
        add_run(p2, str(item.get('label', '')), size=9, color=WHITE)

    # 표 다음 여백
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)

def callout(doc, label, text, top_color_hex=BLUE_HEX):
    """콜아웃 박스: GRAY_LIGHT 배경 + 상단 BLUE 굵은 라인"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(17)
    _shade_cell(cell, GRAY_LIGHT)

    # 셀 테두리: 상단만 굵은 색, 나머지는 옅음
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '24')
    top.set(qn('w:color'), top_color_hex)
    tc_borders.append(top)
    for side in ('left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), GRAY_BORDER)
        tc_borders.append(b)
    tc_pr.append(tc_borders)

    # 라벨
    p_label = cell.paragraphs[0]
    p_label.paragraph_format.space_before = Pt(8)
    p_label.paragraph_format.space_after = Pt(4)
    add_run(p_label, label.upper(), size=8, bold=True, color=CYAN, char_spacing=40)

    # 본문
    p_text = cell.add_paragraph()
    p_text.paragraph_format.space_after = Pt(8)
    p_text.paragraph_format.line_spacing = 1.3
    add_run(p_text, text, size=11, color=NAVY)

    # 표 다음 여백
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

def data_table(doc, headers, rows):
    """데이터 테이블: NAVY 헤더, 교차 GRAY_LIGHT 행"""
    if not headers or not rows:
        return

    n_cols = len(headers)
    table = doc.add_table(rows=len(rows) + 1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # 헤더 행
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        _shade_cell(cell, NAVY_HEX)
        _set_cell_borders(cell, color=GRAY_BORDER, size="4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        add_run(p, str(header), size=10, bold=True, color=WHITE)

    # 본문 행
    for row_idx, row in enumerate(rows, start=1):
        for col_idx in range(n_cols):
            cell = table.cell(row_idx, col_idx)
            _set_cell_borders(cell, color=GRAY_BORDER, size="4")
            # 교차 행 음영
            if row_idx % 2 == 0:
                _shade_cell(cell, GRAY_LIGHT)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            value = row[col_idx] if col_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, str(value), size=10, color=BLACK)

    # 표 다음 여백
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


# ================ 표지 페이지 ================
def add_cover_page(doc, cover):
    """표지: CYAN 라인 → 제목 → 메타 정보 (상단 검은 바 제거됨)"""
    # CYAN 라인 (상단 NAVY/검은 바는 디자인에서 제거)
    p_cyan = doc.add_paragraph()
    p_cyan.paragraph_format.space_before = Pt(2)
    p_cyan.paragraph_format.space_after = Pt(0)
    _bottom_border_paragraph(p_cyan, CYAN_HEX, size="14")

    # 큰 여백
    for _ in range(6):
        doc.add_paragraph()

    # 이브로우
    p_eb = doc.add_paragraph()
    p_eb.paragraph_format.space_after = Pt(8)
    add_run(p_eb, "STRATEGIC BRIEFING | CONSULTING DIAGNOSIS",
            size=8, bold=True, color=CYAN, char_spacing=40)

    # 메인 제목
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_after = Pt(6)
    add_run(p_title, "컨설팅 진단 보고서", size=32, bold=True, color=NAVY)

    # 부제 (고객사명)
    company = cover.get('company_name', '')
    if company:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_after = Pt(8)
        add_run(p_sub, company, size=14, color=NAVY)

    # 보조 부제
    sub = cover.get('subtitle', '')
    if sub:
        p_sub2 = doc.add_paragraph()
        p_sub2.paragraph_format.space_after = Pt(12)
        add_run(p_sub2, sub, size=12, color=BLUE)

    # BLUE divider
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(8)
    p_div.paragraph_format.space_after = Pt(12)
    _bottom_border_paragraph(p_div, BLUE_HEX, size="8")

    # 메타 스택
    meta_items = [
        ("Prepared For", cover.get('prepared_for', '-')),
        ("Prepared By", cover.get('prepared_by', 'venturecompany')),
        ("Date", cover.get('date', '-')),
    ]
    for label, value in meta_items:
        p_m = doc.add_paragraph()
        p_m.paragraph_format.space_after = Pt(2)
        add_run(p_m, f"{label}    ", size=8, bold=True, color=CYAN, char_spacing=40)
        add_run(p_m, str(value), size=10, color=NAVY)

    # 페이지 나눔
    doc.add_page_break()


# ================ 헤더 / 푸터 ================
def setup_header_footer(doc):
    """헤더: venturecompany | 컨설팅 진단 보고서 + CYAN 라인
       푸터: Confidential | Page X of Y"""
    section = doc.sections[0]
    # 표지 첫 페이지는 헤더/푸터 다르게
    section.different_first_page_header_footer = True

    # ── 일반 페이지 헤더 ──
    header = section.header
    h_para = header.paragraphs[0]
    h_para.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_ALIGN_PARAGRAPH.RIGHT)
    add_run(h_para, "venturecompany", size=9, bold=True, color=NAVY)
    h_para.add_run("\t")
    add_run(h_para, "컨설팅 진단 보고서", size=9, color=BLUE)
    _bottom_border_paragraph(h_para, CYAN_HEX, size="6")

    # ── 일반 페이지 푸터 ──
    footer = section.footer
    f_para = footer.paragraphs[0]
    f_para.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_ALIGN_PARAGRAPH.RIGHT)
    _top_border_paragraph(f_para, "D9D9D9", size="6")
    add_run(f_para, "Confidential | © venturecompany", size=8, color=NAVY)
    f_para.add_run("\t")

    # 페이지 번호 (PAGE 필드)
    run_page = f_para.add_run()
    run_page.font.name = FONT
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = NAVY
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run_page._r.append(fld_begin)
    run_page._r.append(instr)
    run_page._r.append(fld_end)

    add_run(f_para, " of ", size=8, color=NAVY)

    run_total = f_para.add_run()
    run_total.font.name = FONT
    run_total.font.size = Pt(8)
    run_total.font.color.rgb = NAVY
    fld_begin2 = OxmlElement('w:fldChar')
    fld_begin2.set(qn('w:fldCharType'), 'begin')
    instr2 = OxmlElement('w:instrText')
    instr2.text = 'NUMPAGES'
    fld_end2 = OxmlElement('w:fldChar')
    fld_end2.set(qn('w:fldCharType'), 'end')
    run_total._r.append(fld_begin2)
    run_total._r.append(instr2)
    run_total._r.append(fld_end2)

    # ── 첫 페이지(표지) 헤더/푸터는 비움 ──
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


# ================ 페이지 설정 ================
def setup_page(doc):
    """US Letter, 1.5cm 좌우 마진, 2.5cm 상하 마진"""
    section = doc.sections[0]
    section.page_height = Cm(27.94)  # 11 inch
    section.page_width = Cm(21.59)   # 8.5 inch
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)
    section.header_distance = Cm(1.27)
    section.footer_distance = Cm(1.27)


# ================ 컴포넌트 디스패치 ================
def render_component(doc, comp):
    """컴포넌트 타입별 렌더링"""
    ctype = comp.get('type', 'paragraph')

    if ctype == 'table':
        data_table(doc, comp.get('headers', []), comp.get('rows', []))

    elif ctype == 'bullets':
        for item in comp.get('items', []):
            bullet(doc, str(item), level=0)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    elif ctype == 'lead_in_bullets':
        for item in comp.get('items', []):
            if isinstance(item, dict):
                lead_in_bullet(doc, item.get('lead', ''), item.get('rest', ''))
            else:
                bullet(doc, str(item), level=0)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    elif ctype == 'kpi_strip':
        kpi_strip(doc, comp.get('items', []))

    elif ctype == 'callout':
        callout(doc, comp.get('label', 'KEY TAKEAWAY'), comp.get('text', ''))

    elif ctype == 'numbered':
        numbered(doc, comp.get('items', []))
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    elif ctype == 'paragraph':
        body(doc, comp.get('text', ''))

    elif ctype == 'h2':
        h2(doc, comp.get('text', ''))

    elif ctype == 'h3':
        h3(doc, comp.get('text', ''))

    else:
        body(doc, str(comp))


# ================ 섹션 렌더링 ================
def render_section(doc, section_data):
    """섹션: eyebrow → H1 → accent bar → intro → components"""
    if section_data.get('eyebrow'):
        eyebrow(doc, section_data['eyebrow'])
    if section_data.get('title'):
        h1(doc, section_data['title'])
    accent_bar(doc)
    if section_data.get('intro'):
        body(doc, section_data['intro'])

    for comp in section_data.get('components', []):
        render_component(doc, comp)


# ================ 메인 빌더 ================
def build_docx(report_data, output_path):
    """보고서 JSON → .docx"""
    doc = Document()

    # 기본 폰트 (정상 출력 보장용)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    r_fonts = rpr.find(qn('w:rFonts'))
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        rpr.insert(0, r_fonts)
    r_fonts.set(qn('w:eastAsia'), FONT)

    setup_page(doc)
    setup_header_footer(doc)

    # 1. 표지 (표지 끝에서 page_break → 이후 본문은 새 페이지부터 시작)
    cover = report_data.get('cover', {})
    add_cover_page(doc, cover)

    # 2. 본문 블록들 — 각 타이틀이 항상 페이지 상단에서 시작하도록 블록마다 페이지 분리
    def _render_meta(meta_rows):
        eyebrow(doc, "REPORT INFORMATION")
        h1(doc, "보고서 정보")
        accent_bar(doc)
        data_table(doc, ["항목", "내용"], meta_rows)

    blocks = []  # 렌더 함수 리스트 (순서대로 각각 새 페이지)

    exec_summary = report_data.get('executive_summary')
    if exec_summary:
        blocks.append(lambda: render_executive_summary(doc, exec_summary))

    meta = report_data.get('meta_table', [])
    if meta:
        blocks.append(lambda: _render_meta(meta))

    for section in report_data.get('sections', []):
        blocks.append(lambda s=section: render_section(doc, s))

    closing = report_data.get('closing')
    if closing:
        blocks.append(lambda: render_section(doc, closing))

    # 첫 블록은 표지의 page_break 덕에 이미 새 페이지에서 시작.
    # 두 번째 블록부터는 앞에 page_break 를 넣어 타이틀이 페이지 상단에서 시작하도록 한다.
    for i, render in enumerate(blocks):
        if i > 0:
            doc.add_page_break()
        render()

    doc.save(output_path)
    print(f"✅ .docx 생성 완료: {output_path}")


# ================ EXECUTIVE SUMMARY 전용 렌더링 ================
def render_executive_summary(doc, exec_data):
    """EXECUTIVE SUMMARY: eyebrow → H1 → accent bar → intro → kpi_strip → callout"""
    if not exec_data:
        return
    if exec_data.get('eyebrow'):
        eyebrow(doc, exec_data['eyebrow'])
    if exec_data.get('title'):
        h1(doc, exec_data['title'])
    accent_bar(doc)
    if exec_data.get('intro'):
        body(doc, exec_data['intro'])
    if exec_data.get('kpi'):
        kpi_strip(doc, exec_data['kpi'])
    if exec_data.get('key_takeaway'):
        callout(doc, "KEY TAKEAWAY", exec_data['key_takeaway'])
