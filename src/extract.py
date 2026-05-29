"""업로드된 파일에서 텍스트를 추출."""
from __future__ import annotations

import io


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()

    if name.endswith(".pdf"):
        return _pdf(data)
    if name.endswith(".docx"):
        return _docx(data)
    if name.endswith((".txt", ".md", ".csv", ".json")):
        return _decode_text(data)
    if name.endswith(".doc"):
        raise ValueError(
            f"'{filename}': 구형 .doc 포맷은 지원하지 않습니다. .docx 또는 .pdf 로 저장 후 업로드하세요."
        )

    # 확장자 불명 → 텍스트로 시도
    try:
        return _decode_text(data)
    except ValueError:
        raise ValueError(f"'{filename}': 텍스트를 추출할 수 없는 파일 형식입니다.")


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "cp949", "euc-kr"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("텍스트 인코딩을 인식할 수 없습니다.")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def _docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    lines = [p.text for p in doc.paragraphs]
    # 표 안의 텍스트도 포함
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines).strip()
